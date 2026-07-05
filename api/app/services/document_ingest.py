"""Extract plain text / markdown from uploaded itinerary documents.

Supports Excel (.xlsx), PDF (.pdf), and Word (.docx). The extracted text is fed
to the AI structuring pass, so the goal is a faithful, readable rendering of the
document's content (tables become markdown tables).
"""

from __future__ import annotations

import io

from fastapi import HTTPException, status

_MAX_CHARS = 200_000


def extract_text(filename: str, data: bytes) -> str:
    """Return extracted text for the given uploaded file.

    Raises HTTPException(415) for unsupported types.
    """
    name = (filename or "").lower()

    if name.endswith(".xlsx"):
        text = _extract_xlsx(data)
    elif name.endswith(".pdf"):
        text = _extract_pdf(data)
    elif name.endswith(".docx"):
        text = _extract_docx(data)
    else:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail="Unsupported file type. Upload an .xlsx, .pdf, or .docx file.",
        )

    text = text.strip()
    if not text:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Could not extract any text from the uploaded file.",
        )
    return text[:_MAX_CHARS]


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    blocks: list[str] = []
    for ws in wb.worksheets:
        rows = [
            [("" if c is None else str(c)).strip() for c in row]
            for row in ws.iter_rows(values_only=True)
        ]
        # Drop fully-empty rows.
        rows = [r for r in rows if any(cell for cell in r)]
        if not rows:
            continue
        blocks.append(f"## Sheet: {ws.title}")
        blocks.append(_rows_to_markdown(rows))
    wb.close()
    return "\n\n".join(blocks)


def _rows_to_markdown(rows: list[list[str]]) -> str:
    width = max(len(r) for r in rows)
    padded = [r + [""] * (width - len(r)) for r in rows]
    header = padded[0]
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * width) + " |"]
    for r in padded[1:]:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            parts.append(f"## Page {i}\n{text.strip()}")
    return "\n\n".join(parts)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        rows = [r for r in rows if any(c for c in r)]
        if rows:
            parts.append(_rows_to_markdown(rows))
    return "\n\n".join(parts)
